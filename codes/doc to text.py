from docx import *
import re
import json
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

#_-----------for document(excluding tables) to text-----------
def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

#_-----------for tables to text-----------
def table_print(block):
    table=block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text,'  ',end='')
                #y.write(paragraph.text)
                #y.write('  ')
        print("\n")            
    

document = Document('C:/Users/AMAN SHAKYA/Desktop/New Microsoft Word Document.docx')
for block in iter_block_items(document):
    if isinstance(block, Paragraph):
        print(block.text)
    elif isinstance(block, Table):
        table_print(block)

#_-----------Declare Variables-----------
bolds=[]
emails=[]
phones=[]

#-----------Extract Elements From the Word File-----------
for para in document.paragraphs:

    #Find email and phone numbers within the paragraph text
    text = para.text
    email_list = re.findall(r'[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+',text)
    phone_list=re.findall(r'[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]',text)

    for email in email_list:
        emails.append(email)

    for phone in phone_list:
        phones.append(phone)

    #Find the bold style within the word document
    for run in para.runs:
        if run.bold :
            bolds.append(run.text)

#-----------Create Output-----------
style_Dict={'emails':emails,'phone_numbers':phones,'bold_phrases':bolds}
              

print("\nWord File Output:\n")

r = json.dumps(style_Dict)
loaded_r = json.loads(r)
print("\n",json.dumps(loaded_r,indent=4, sort_keys=False))  #Pretty print the JSON output

